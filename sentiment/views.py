from django.shortcuts import render
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing.sequence import pad_sequences
from tensorflow.keras.preprocessing.text import tokenizer_from_json
from django.http import FileResponse, HttpResponseRedirect
from django.urls import reverse
from django.core.files.storage import default_storage
from django.conf import settings
import os
import json
import os

# Загрузка модели
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
model_path = os.path.join(BASE_DIR, 'src/Sentiment Model.keras')
model = load_model(model_path)

# Загрузка токенизатора
tokenizer_path = os.path.join(BASE_DIR, 'src/tokenizer.json')
with open(tokenizer_path, 'r') as f:
    tokenizer_data = json.load(f)
tokenizer = tokenizer_from_json(tokenizer_data)

# Параметры обработки текста
maxlen = 150

def index(request):
    result = None
    if request.method == 'POST':
        text = request.POST.get('comment')  # Получаем текст от пользователя
        if text:
            # Токенизация текста
            tokenized_text = tokenizer.texts_to_sequences([text])
            # Паддинг
            padded_text = pad_sequences(tokenized_text, maxlen=maxlen)

            # Предсказание
            prediction = model.predict(padded_text)
            sentiment = {0: 'Нейтральный', 1: 'Позитивный', 2: 'Негативный'}
            result = sentiment[prediction.argmax()]  # Определяем тональность

    return render(request, 'sentiment/index.html', {'result': result})


def create_report(csv_path, docx_path):
    from docx import Document
    doc = Document()
    doc.add_heading('Отчет по анализу тональности', 0)
    doc.add_paragraph('Здесь будет содержаться результат анализа файла CSV.')
    doc.save(docx_path)

def upload_csv(request):
    if request.method == 'POST' and request.FILES.get('csv_file'):
        csv_file = request.FILES['csv_file']
        filename = default_storage.save('temp_uploaded.csv', csv_file)
        csv_path = os.path.join(settings.MEDIA_ROOT, filename)

        report_path = os.path.join(settings.MEDIA_ROOT, 'report.docx')

        create_report(csv_path, report_path)

        report_url = settings.MEDIA_URL + 'report.docx'
        return render(request, 'sentiment/index.html', {
            'report_generated': True,
            'report_url': report_url
        })

    return HttpResponseRedirect(reverse('index'))
