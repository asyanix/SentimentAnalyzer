import nltk
from nltk.corpus import stopwords

def download_stopwords():
    try:
        _ = stopwords.words('russian')
    except LookupError:
        nltk.download('stopwords')

download_stopwords()  
RUSSIAN_STOPWORDS = stopwords.words('russian')

from tensorflow.keras.preprocessing.sequence import pad_sequences
from sklearn.feature_extraction.text import CountVectorizer
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor, Inches
from django.conf import settings
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
import concurrent.futures
from tqdm import tqdm
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg') 
import re
import os
import concurrent.futures

def clean_text(text):
    text = re.sub(r'[^\w\s]', '', text)   
    text = re.sub(r'\d+', '', text)       
    return text.lower().strip()

def predict_batch(batch_texts, tokenizer, model, maxlen):
    sequences = tokenizer.texts_to_sequences(batch_texts)
    padded = pad_sequences(sequences, maxlen=maxlen)
    preds = model.predict(padded)
    return preds.argmax(axis=1)

def threaded_predict(texts, tokenizer, model, maxlen=150, batch_size=512, max_workers=4):
    batches = [texts[i:i + batch_size] for i in range(0, len(texts), batch_size)]
    results = [None] * len(batches)

    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_idx = {
            executor.submit(predict_batch, batch, tokenizer, model, maxlen): idx
            for idx, batch in enumerate(batches)
        }
        for future in concurrent.futures.as_completed(future_to_idx):
            idx = future_to_idx[future]
            try:
                results[idx] = future.result()
            except Exception as e:
                print(f"Ошибка в батче {idx}: {e}")
                results[idx] = np.zeros(len(batches[idx]), dtype=int)

    return np.concatenate(results)

def analyze_csv(csv_path, model, tokenizer, maxlen=150):
    df = pd.read_csv(csv_path, sep=',', engine='python', encoding='utf-8', usecols=[0], header=0)
    df.columns = ['text']  

    if 'text' not in df.columns:
        raise ValueError("CSV должен содержать столбец 'text' с комментариями.")

    df['text'] = df['text'].astype(str).apply(clean_text)
    texts = df['text'].tolist()

    labels = threaded_predict(texts, tokenizer, model, maxlen=maxlen, batch_size=512, max_workers=4)

    df['label_code'] = labels
    return df
    
def enter(doc, count):
    for i in range(count):
        doc.add_paragraph()

def par(doc, text, align = WD_ALIGN_PARAGRAPH.JUSTIFY, is_bold = False, is_italic = False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = is_bold
    run.italic = is_italic
    return p

def extract_top_ngrams(texts, ngram_range=(1, 1), top_n=5):
    if not texts:
        return []
    vectorizer = CountVectorizer(stop_words=RUSSIAN_STOPWORDS, ngram_range=ngram_range)
    X = vectorizer.fit_transform(texts)
    freqs = np.asarray(X.sum(axis=0)).ravel()
    vocab = vectorizer.get_feature_names_out()
    return sorted(zip(vocab, freqs), key=lambda x: x[1], reverse=True)[:top_n]

def add_ngrams_table(doc, ngrams, title):
    table = doc.add_table(rows=len(ngrams) + 1, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = title
    table.cell(0, 1).text = 'Частота'
    for i, (text, freq) in enumerate(ngrams, start=1):
        table.cell(i, 0).text = text
        table.cell(i, 1).text = str(freq)

def add_top_ngrams_section(doc, df):
    label_map = {0: 'Нейтральный', 1: 'Позитивный', 2: 'Негативный'}
    index = 6

    for label_code, label_name in label_map.items():
        texts = df[df['label_code'] == label_code]['text'].astype(str).tolist()
        if not texts:
            continue

        par(doc, f'{label_name} класс:')

        # Топ униграммы
        top_unigrams = extract_top_ngrams(texts, ngram_range=(1, 1), top_n=5)
        add_ngrams_table(doc, top_unigrams, 'Топ-слово')
        par(doc, f'Таблица {index}. Топ-слова для класса «{label_name}»', WD_ALIGN_PARAGRAPH.CENTER)
        enter(doc, 1)
        index += 1

        # Топ биграммы
        top_bigrams = extract_top_ngrams(texts, ngram_range=(2, 2), top_n=5)
        add_ngrams_table(doc, top_bigrams, 'Топ-биграмма')
        par(doc, f'Таблица {index}. Топ-биграммы для класса «{label_name}»', WD_ALIGN_PARAGRAPH.CENTER)
        enter(doc, 1)
        index += 1

def create_report(df, report_path):
    total = len(df)
    label_map = {0: 'Нейтральный', 1: 'Позитивный', 2: 'Негативный'}
    labels = ['Нейтральный', 'Позитивный', 'Негативный']
    label_counts = df['label_code'].value_counts().sort_index()
    label_percent = label_counts / total * 100
    positive_count = label_counts.get(1, 0)
    neutral_count = label_counts.get(0, 0)
    negative_count = label_counts.get(2, 0)
    
    doc = Document()
    
    # ТИТУЛЬНИК
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)  
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    
    enter(doc, 5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Отчёт об анализе тональности больших данных')
    run.bold = True
    run.font.size = Pt(24)

    enter(doc, 15)    
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = title.add_run("Является продуктом нейросетевой модели Sentiment Model\n")
    info.font.size = Pt(16)
    date = title.add_run(f"{datetime.today().strftime('%d.%m.%Y')}")
    date.italic = True
    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # ЦЕЛИ И ЗАДАЧИ
    
    par(doc, 'ЦЕЛИ И ЗАДАЧИ', WD_ALIGN_PARAGRAPH.CENTER, True)
    
    goal = par(doc, '\tЦелью ', is_bold= True)
    run = goal.add_run('настоящего исследования является проведение количественно-качественного анализа субъективной текстовой информации с использованием методов автоматического распознавания тональности. ')
    
    tasks = par(doc, '\tДля достижения поставленной цели были сформулированы следующие ')
    run = tasks.add_run('задачи:')
    run.bold = True

    par(doc, '\t1. Описание выборки текстов;')
    par(doc, '\t2. Предобработка текстовых данных;')
    par(doc, '\t3. Автоматическая классификация эмоциональной окраски текстов;')
    par(doc, '\t4. Анализ распределения тональностей;')
    par(doc, '\t5. Построение визуальных и статистических представлений;')
    par(doc, '\t6. Формулировка выводов и рекомендаций.') 
    enter(doc, 1)
    
    # ОПИСАНИЕ ДАННЫХ
    
    par(doc, 'ОПИСАНИЕ ДАННЫХ', WD_ALIGN_PARAGRAPH.CENTER, True)
    par(doc, f'\tОбъём (общее число текстовых единиц): {total}.')
    par(doc, '\tЯзык файла: Русский.')
    par(doc, '\tИсточник: файл формата csv.')
    par(doc, '\tКачество: текстовая нормализация, векторизация и токенизация текста.')
    enter(doc, 1)    
    
    # МЕТОДОЛОГИЯ АНАЛИЗА

    par(doc, 'МЕТОДОЛОГИЯ АНАЛИЗА', WD_ALIGN_PARAGRAPH.CENTER, True)
    par(doc, '\tМодель: Трехслойная NLP-модель (Embedding, GRU, Dense) с оптимизатором AdamW и технологией EarlyStopping.')
    par(doc, '\tПороговые значения: «Позитивный», «Нейтральный» и «Негативный».')
    enter(doc, 1)    
    
    # АНАЛИЗ БОЛЬШИХ ДАННЫХ
    
    par(doc, 'АНАЛИЗ БОЛЬШИХ ДАННЫХ', WD_ALIGN_PARAGRAPH.CENTER, True)
    par(doc, '\t1. Распределение тональностей', is_italic=True)
    par(doc, '\tКоличество текстов в каждой категории: позитивный, нейтральный, негативный.')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'  
    table.cell(0, 0).text = 'Классы тональности'
    table.cell(0, 1).text = 'Количество текстов'
    table.cell(1, 0).text = 'Позитивный'
    table.cell(1, 1).text = f'{positive_count}'
    table.cell(2, 0).text = 'Нейтральный'
    table.cell(2, 1).text = f'{neutral_count}'
    table.cell(3, 0).text = 'Негативный'
    table.cell(3, 1).text = f'{negative_count}'
    par(doc, 'Таблица 1. Количество текстов, классифицированных по трем категориям тональности.', WD_ALIGN_PARAGRAPH.CENTER)
    enter(doc, 1)    
    
    # Гистограмма
    values = [neutral_count, positive_count, negative_count]
    plt.figure(figsize=(6, 4))
    bars = plt.bar(labels, values, color=['#1f77b4', '#2ca02c', '#d62728'])
    plt.title('Распределение тональностей')
    plt.xlabel('Классы тональности')
    plt.ylabel('Количество текстов')
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    image_path = os.path.join(settings.MEDIA_ROOT, 'histogram.png')
    plt.savefig(image_path)
    plt.close()
    
    par(doc, 'Визуализируем результаты анализа с помощью гистограммы.')
    enter(doc, 1)
    doc.add_picture(image_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    par(doc, 'Рисунок 1. Гистограмма, отражающая абсолютное количество текстов в каждой категории тональности.', WD_ALIGN_PARAGRAPH.CENTER)
    
    enter(doc, 2)
    
    par(doc, '2. Баланс классов', is_italic=True)
    par(doc, 'Соотношение между количеством текстов разных тональностей.')
    
    positive_percent = round(label_percent.get(1, 0), 2)
    neutral_percent = round(label_percent.get(0, 0), 2)
    negative_percent = round(label_percent.get(2, 0), 2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'  
    table.cell(0, 0).text = 'Классы тональности'
    table.cell(0, 1).text = 'Доля (в процентах)'
    table.cell(1, 0).text = 'Позитивный'
    table.cell(1, 1).text = f'{positive_percent}'
    table.cell(2, 0).text = 'Нейтральный'
    table.cell(2, 1).text = f'{neutral_percent}'
    table.cell(3, 0).text = 'Негативный'
    table.cell(3, 1).text = f'{negative_percent}'
    par(doc, 'Таблица 2. Доля каждой категории тональности в общем объеме проанализированных текстов.', WD_ALIGN_PARAGRAPH.CENTER)
    enter(doc, 1)
    
    # Круговая диаграмма
    plt.figure(figsize=(6, 4))
    sizes = [neutral_percent, positive_percent, negative_percent]
    colors = ['#1f77b4', '#2ca02c', '#d62728']

    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=colors)
    plt.axis('equal')
    plt.title('Соотношение классов тональности')
    pie_chart_path = os.path.join(settings.MEDIA_ROOT, 'pie_chart.png')
    plt.savefig(pie_chart_path)
    plt.close()
    
    par(doc, 'Визуализируем результаты анализа с помощью круговой диаграммы.')
    enter(doc, 1)
    doc.add_picture(pie_chart_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    par(doc, 'Рисунок 2. Круговая диаграмма, демонстрирующая относительное распределение классов тональности.', WD_ALIGN_PARAGRAPH.CENTER)
    
    enter(doc, 1)
    
    par(doc, '3. Сентимент-индекс', is_italic=True)
    par(doc, 'Сводная оценка "эмоционального фона" всей выборки.')
    doc.add_picture('sentiment/sentiment_index.png', width=Inches(3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    par(doc, 'Рисунок 3. Формула сентимент-индекса', WD_ALIGN_PARAGRAPH.CENTER)
    enter(doc, 1)
    
    sentiment_index = (positive_count - negative_count) / total
    result = f"Сентимент-индекс равен {sentiment_index:.3f}, следовательно можно сделать вывод, что в объеме проанализированных текстов "
    if sentiment_index == 0: 
        result += "полная нейтральность"
    elif sentiment_index > 0: 
        result += "доминирует позитивная окраска"
    else: 
        result += "доминирует отрицательная окраска"
    
    par(doc, result)
    
    enter(doc, 1)
    
    par(doc, '4. Распределение по длине текста', is_italic=True)
    par(doc, 'Анализ длины текстов в разных тональностях.')
    
    df['text_length'] = df['text'].apply(lambda x: len(str(x).split()))
    grouped = df.groupby('label_code')['text_length']

    avg_all = df['text_length'].mean()
    avg_pos = grouped.mean().get(1, 0)
    avg_neg = grouped.mean().get(2, 0)
    avg_neu = grouped.mean().get(0, 0)

    min_all = df['text_length'].min()
    min_pos = grouped.min().get(1, 0) if grouped.min().get(1, 0) != 0 else 1
    min_neg = grouped.min().get(2, 0) if grouped.min().get(2, 0) != 0 else 1
    min_neu = grouped.min().get(0, 0) if grouped.min().get(0, 0) != 0 else 1

    max_all = df['text_length'].max()
    max_pos = grouped.max().get(1, 0)
    max_neg = grouped.max().get(2, 0)
    max_neu = grouped.max().get(0, 0)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'  
    table.cell(0, 0).text = 'Показатель'
    table.cell(0, 1).text = 'Количество слов'
    table.cell(1, 0).text = 'Средняя длина текста'
    table.cell(1, 1).text = f'{avg_all:.3f}'
    table.cell(2, 0).text = 'Средняя длина положительного текста'
    table.cell(2, 1).text = f'{avg_pos:.3f}'
    table.cell(3, 0).text = 'Средняя длина отрицательного текста'
    table.cell(3, 1).text = f'{avg_neg:.3f}'
    table.cell(4, 0).text = 'Средняя длина нейтрального текста'
    table.cell(4, 1).text = f'{avg_neu:.3f}'
    par(doc, 'Таблица 3. Средняя длина текста для каждого класса тональности', WD_ALIGN_PARAGRAPH.CENTER)
    enter(doc, 1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'  
    table.cell(0, 0).text = 'Показатель'
    table.cell(0, 1).text = 'Количество слов'
    table.cell(1, 0).text = 'Минимальная длина текста'
    table.cell(1, 1).text = f'{min_all}'
    table.cell(2, 0).text = 'Минимальная длина положительного текста'
    table.cell(2, 1).text = f'{min_pos}'
    table.cell(3, 0).text = 'Минимальная длина отрицательного текста'
    table.cell(3, 1).text = f'{min_neg}'
    table.cell(4, 0).text = 'Минимальная длина нейтрального текста'
    table.cell(4, 1).text = f'{min_neu}'
    par(doc, 'Таблица 4. Минимальная длина текста для класса тональности', WD_ALIGN_PARAGRAPH.CENTER)
    enter(doc, 1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'  
    table.cell(0, 0).text = 'Показатель'
    table.cell(0, 1).text = 'Количество слов'  
    table.cell(1, 0).text = 'Максимальная длина текста'
    table.cell(1, 1).text = f'{max_all}'
    table.cell(2, 0).text = 'Максимальная длина положительного текста'
    table.cell(2, 1).text = f'{max_pos}'
    table.cell(3, 0).text = 'Максимальная длина отрицательного текста'
    table.cell(3, 1).text = f'{max_neg}'
    table.cell(4, 0).text = 'Максимальная длина нейтрального текста'
    table.cell(4, 1).text = f'{max_neu}'
    par(doc, 'Таблица 5. Максимальная длина текста для класса тональности', WD_ALIGN_PARAGRAPH.CENTER)
    enter(doc, 1)
    
    # Ящик с усами
    df['label_text'] = df['label_code'].map(label_map)
    df.boxplot(column='text_length', by='label_text', grid=False, patch_artist=True,
            boxprops=dict(facecolor='lightblue', color='blue'),
            medianprops=dict(color='red'),
            whiskerprops=dict(color='gray'),
            capprops=dict(color='gray'))

    plt.title('Распределение длины текстов по тональности')
    plt.suptitle('')
    plt.xlabel('Класс тональности')
    plt.ylabel('Длина текста (в словах)')
    plt.grid(True, linestyle='--', alpha=0.7)
    box_path = os.path.join(settings.MEDIA_ROOT, 'box_plot.png')
    plt.tight_layout()
    plt.savefig(box_path)
    plt.close()

    par(doc, 'Визуализируем результаты анализа с помощью диаграммы "ящик с усами".')
    enter(doc, 1)
    doc.add_picture(box_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    par(doc, 'Рисунок 4. Ящик с усами', WD_ALIGN_PARAGRAPH.CENTER)
    
    enter(doc, 1)
    
    par(doc, '5. Топ-слова и n-граммы по тональности', is_italic=True)
    par(doc, 'Наиболее часто встречающиеся слова или словосочетания в каждом классе.')
    add_top_ngrams_section(doc, df)
    
    doc.save(report_path)
